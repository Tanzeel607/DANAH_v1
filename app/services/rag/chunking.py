"""Text extraction and paragraph-aware chunking.

Chunks target `CHUNK_SIZE_TOKENS` (~800) with `CHUNK_OVERLAP_TOKENS` (~150) of overlap and are
split on paragraph boundaries wherever possible. That boundary choice is not cosmetic: a citation
is shown to an executive, so a chunk that begins mid-sentence produces a quotation that reads as
nonsense and erodes trust in the answer.

Token counting uses `tiktoken` (cl100k_base) as a provider-neutral estimator — Anthropic ships no
local tokenizer, and chunk sizing is not a correctness boundary (docs/DECISIONS.md #13).
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from functools import lru_cache
from io import BytesIO
from pathlib import Path

import structlog
import tiktoken

from app.exceptions import RetrievalError

log = structlog.get_logger(__name__)

# Paragraph break: a blank line, tolerating trailing whitespace.
_PARAGRAPH_RE = re.compile(r"\n\s*\n")
# Sentence end, used only when a single paragraph is itself larger than a chunk.
_SENTENCE_RE = re.compile(r"(?<=[.!?؟।])\s+")
_WHITESPACE_RE = re.compile(r"[ \t]+")


@lru_cache(maxsize=1)
def _encoder() -> tiktoken.Encoding:
    return tiktoken.get_encoding("cl100k_base")


def count_tokens(text: str) -> int:
    return len(_encoder().encode(text, disallowed_special=()))


@dataclass(slots=True)
class Chunk:
    index: int
    content: str
    token_count: int


# ---------------------------------------------------------------------------
# Extraction
# ---------------------------------------------------------------------------
def extract_text(data: bytes, *, filename: str, mime_type: str = "") -> str:
    """Extract plain text from pdf / docx / txt / md / html.

    The extension drives the choice (the browser-supplied mime type is not trustworthy).
    """
    suffix = Path(filename).suffix.lower().lstrip(".")

    if suffix == "pdf":
        return _extract_pdf(data)
    if suffix == "docx":
        return _extract_docx(data)
    if suffix in {"html", "htm"}:
        return _extract_html(data)
    if suffix in {"txt", "md", "markdown", "text"}:
        return _decode(data)

    raise RetrievalError(
        f"Cannot extract text from '.{suffix}' files.",
        code="unsupported_file_type",
        detail={"filename": filename, "mime_type": mime_type},
    )


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1256", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # cp1256 is the common Arabic Windows codepage; latin-1 never fails, so this is unreachable
    # in practice and exists only to keep the type checker honest.
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except (PdfReadError, ValueError, OSError) as exc:
        raise RetrievalError(
            "The PDF could not be read. It may be corrupt, encrypted or image-only.",
            code="pdf_extraction_failed",
            detail={"error": str(exc)},
        ) from exc

    text = "\n\n".join(p.strip() for p in pages if p.strip())
    if not text.strip():
        # A scanned PDF yields zero characters. Saying so beats indexing an empty document and
        # letting the user wonder why chat cannot find anything in it.
        raise RetrievalError(
            "No text could be extracted from this PDF. If it is a scan, it needs OCR first.",
            code="pdf_no_text",
        )
    return text


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_html(data: bytes) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(_decode(data), "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [line.strip() for line in text.splitlines()]
    return "\n\n".join(line for line in lines if line)


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------
def normalise(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WHITESPACE_RE.sub(" ", text)
    # Collapse 3+ newlines to exactly the paragraph separator.
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def chunk_text(text: str, *, chunk_size: int, overlap: int) -> list[Chunk]:
    """Split into ~`chunk_size`-token chunks with `overlap` tokens of context carried forward.

    Paragraphs are the unit of assembly. A paragraph larger than a whole chunk is split on
    sentence boundaries; a single sentence larger than a chunk is split on tokens as a last
    resort (it is almost always a table or a wall of code).
    """
    if chunk_size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size")

    cleaned = normalise(text)
    if not cleaned:
        return []

    units = _split_into_units(cleaned, chunk_size)

    chunks: list[Chunk] = []
    current: list[tuple[str, int]] = []
    current_tokens = 0

    for unit, unit_tokens in units:
        if current and current_tokens + unit_tokens > chunk_size:
            chunks.append(_emit(chunks, current))
            current, current_tokens = _carry_overlap(current, overlap)

        current.append((unit, unit_tokens))
        current_tokens += unit_tokens

    if current:
        chunks.append(_emit(chunks, current))

    return chunks


def _emit(existing: list[Chunk], parts: list[tuple[str, int]]) -> Chunk:
    content = "\n\n".join(p for p, _ in parts).strip()
    return Chunk(index=len(existing), content=content, token_count=count_tokens(content))


def _carry_overlap(parts: list[tuple[str, int]], overlap: int) -> tuple[list[tuple[str, int]], int]:
    """Keep whole trailing units totalling up to `overlap` tokens as the next chunk's preamble.

    Overlapping by whole units (not by a raw token count) is what keeps the carried-over text
    readable — the next chunk starts at a paragraph boundary, not mid-word.
    """
    if overlap <= 0:
        return [], 0

    carried: list[tuple[str, int]] = []
    total = 0
    for unit, tokens in reversed(parts):
        if total + tokens > overlap:
            break
        carried.insert(0, (unit, tokens))
        total += tokens
    return carried, total


def _split_into_units(text: str, chunk_size: int) -> list[tuple[str, int]]:
    """Paragraphs, subdivided until every unit fits inside one chunk."""
    units: list[tuple[str, int]] = []

    for paragraph in _PARAGRAPH_RE.split(text):
        para = paragraph.strip()
        if not para:
            continue

        tokens = count_tokens(para)
        if tokens <= chunk_size:
            units.append((para, tokens))
            continue

        for sentence in _split_sentences(para, chunk_size):
            units.append((sentence, count_tokens(sentence)))

    return units


def _split_sentences(paragraph: str, chunk_size: int) -> list[str]:
    pieces: list[str] = []
    buffer: list[str] = []
    buffer_tokens = 0

    for sentence in _SENTENCE_RE.split(paragraph):
        sentence = sentence.strip()
        if not sentence:
            continue
        tokens = count_tokens(sentence)

        if tokens > chunk_size:
            if buffer:
                pieces.append(" ".join(buffer))
                buffer, buffer_tokens = [], 0
            pieces.extend(_split_by_tokens(sentence, chunk_size))
            continue

        if buffer_tokens + tokens > chunk_size:
            pieces.append(" ".join(buffer))
            buffer, buffer_tokens = [], 0

        buffer.append(sentence)
        buffer_tokens += tokens

    if buffer:
        pieces.append(" ".join(buffer))
    return pieces


def _split_by_tokens(text: str, chunk_size: int) -> list[str]:
    """Last resort for a single 'sentence' longer than a chunk (tables, code, minified text)."""
    enc = _encoder()
    token_ids = enc.encode(text, disallowed_special=())
    return [enc.decode(token_ids[i : i + chunk_size]) for i in range(0, len(token_ids), chunk_size)]


def detect_language(text: str) -> str:
    """'ar' if the text is substantially Arabic, else 'en'.

    Counting Arabic-script codepoints is enough here: the corpus is EN/AR by product definition,
    and this only picks the FTS/answer language hint — a wrong guess degrades ranking slightly,
    it does not break retrieval.
    """
    sample = text[:4000]
    letters = [c for c in sample if c.isalpha()]
    if not letters:
        return "en"
    arabic = sum(1 for c in letters if "؀" <= c <= "ۿ" or "ݐ" <= c <= "ݿ")
    return "ar" if arabic / len(letters) > 0.30 else "en"
